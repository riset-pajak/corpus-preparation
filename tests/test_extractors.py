"""Unit test: extractors dan pipeline."""

import io
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

from corpusprep.extractors import extract_text_from_pdf, extract_text_from_docx


def test_extract_text_from_docx(tmp_path):
    """Test ekstraksi DOCX."""
    from docx import Document

    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Pasal 1")
    doc.add_paragraph("Setiap wajib pajak wajib membayar pajak.")
    doc.add_paragraph("Pasal 2")
    doc.add_paragraph("Ketentuan lebih lanjut diatur oleh Menteri Keuangan.")
    doc.save(doc_path)

    result = extract_text_from_docx(doc_path)
    assert "Pasal 1" in result
    assert "wajib pajak" in result
    assert "Pasal 2" in result


def test_extract_text_from_pdf_pdfminer(tmp_path):
    """Test ekstraksi PDF dengan reportlab-made PDF."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed, using mock instead")
        return

    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "Pasal 1")
    c.drawString(100, 730, "Setiap wajib pajak wajib membayar pajak.")
    c.drawString(100, 710, "Pasal 2")
    c.drawString(100, 690, "Ketentuan lebih lanjut diatur.")
    c.save()

    result = extract_text_from_pdf(pdf_path)
    assert "Pasal 1" in result
    assert "wajib pajak" in result


def test_extract_text_from_pdf_mock(tmp_path):
    """Test ekstraksi PDF dengan mock (tanpa reportlab)."""
    pdf_path = tmp_path / "dummy.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 dummy")

    with patch("pdfplumber.open") as mock_open:
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Pasal 1\nSetiap wajib pajak wajib membayar pajak.\n\nPasal 2\nKetentuan lebih lanjut diatur."
        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_open.return_value.__enter__ = MagicMock(return_value=mock_pdf)
        mock_open.return_value.__exit__ = MagicMock(return_value=False)

        result = extract_text_from_pdf(pdf_path)
        assert "Pasal 1" in result
        assert "wajib pajak" in result

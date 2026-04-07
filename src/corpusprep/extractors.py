"""Ekstraksi teks dari berbagai format dokumen."""

from __future__ import annotations

from pathlib import Path


def extract_text_from_pdf(file_path: Path) -> str:
    """Ekstrak teks dari PDF menggunakan pdfplumber."""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def extract_text_from_docx(file_path: Path) -> str:
    """Ekstrak teks dari DOCX menggunakan python-docx."""
    from docx import Document

    doc = Document(file_path)
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # Tabel juga
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                paragraphs.append(line)

    return "\n\n".join(paragraphs)
